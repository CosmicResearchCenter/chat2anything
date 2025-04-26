from docx import Document
from core.ocr.ocr_model import OCR_Model
from langchain_core.documents import Document as LangChainDocument
from core.rag.splitter.structured_file import TextSplitter
from config.config_info import settings
from config.splitter_model import SplitterModel
from typing import List
import io
from PIL import Image


class DocxSplitter(TextSplitter):
    def __init__(self, file_path: str,splitter_args=None, splitter_model: SplitterModel = settings.SPPLITTER_MODEL, *args, **kwargs):
        super().__init__(splitter_args=splitter_args,SPPLITTER_MODEL=splitter_model,*args, **kwargs)
        self.ocr_model = OCR_Model()
        self.file_path = file_path
        # self.splitter_pattern = splitter_pattern
        self.splitter_model = splitter_model

    def load(self):
        # 读取 docx 文件
        doc = Document(self.file_path)
        full_text = ""
        
        # 处理段落文本
        for para in doc.paragraphs:
            full_text += para.text + "\n"
        
        # 处理表格
        for table in doc.tables:
            # 获取表格行数与列数
            num_cols = len(table.columns)
            # 创建表头分隔行
            full_text += "|" + " | ".join(["---"] * num_cols) + "|\n"
            
            for row in table.rows:
                # 逐行读取单元格内容并添加分隔符
                row_text = "| " + " | ".join(cell.text.strip() for cell in row.cells) + " |\n"
                full_text += row_text
            full_text += "\n\n"
        
        # 专门处理文档中的所有图片 - 更可靠的方法
        try:
            # 获取所有关系
            rels = doc.part.rels
            for rel in rels.values():
                # 检查关系类型是否为图片
                if "image" in rel.reltype:
                    print(f"找到图片: {rel.target_ref}")
                    try:
                        # 获取图片部分
                        image_part = rel.target_part
                        # 获取图片二进制数据
                        image_bytes = image_part.blob
                        # 使用OCR识别图片文本
                        ocr_text = self.ocr_model.ocr_image_by_image_bytes(image_bytes)
                        if ocr_text:
                            full_text += f"[图片OCR内容]: {ocr_text}\n\n"
                            print(f"成功识别图片文本: {ocr_text[:50]}...")
                        else:
                            print("图片OCR未返回文本")
                    except Exception as e:
                        print(f"处理图片时出错: {e}")
        except Exception as e:
            print(f"处理文档图片时出现错误: {e}")

        print(full_text)
        return full_text
        
    def split(self)->List[LangChainDocument]:
        full_text = self.load()
        return super().split(full_text)

if __name__ == "__main__":
    docx_splitter = DocxSplitter("/Users/markyangkp/Documents/信息整理/常用校园信息集合.docx")
    docx_splitter.load()